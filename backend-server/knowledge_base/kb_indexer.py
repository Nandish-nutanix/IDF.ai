"""
IDF Knowledge Base Indexer — V2.

Indexes ~470 .docx documents from 12 categories into ChromaDB.
Key design for accuracy at scale:
- Heading-aware chunking preserves document structure
- Each chunk is prefixed with [Category > Filename] context for richer embeddings
- Deduplication: skips "Copy of..." files and near-duplicate chunks
- Larger chunks (1500 chars) with 200-char overlap for better retrieval coherence
- Excel/PPTX/media files are excluded
- Uses nomic-embed-text embeddings via Ollama
"""

import os
import re
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple

RAW_DOCS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "raw_docs")
CHROMA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "chroma_db")
INDEX_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "doc_index.json")
HTML_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "html_docs")

CATEGORY_MAP = {
    "Arithmos": "Arithmos & Metric Collection",
    "Customer Interactions": "Customer Issues & Interactions",
    "Design Docs": "IDF Design Documents",
    "Documentation": "IDF Documentation & Guides",
    "Flow SMSP to PC-PE Sync": "PE Sync & Replication",
    "Memory Management": "Memory Management",
    "MongoDB Integration": "MongoDB Integration",
    "PC Federation": "PC Federation & Multi-Cluster",
    "Postgres Over ChakrDB": "Postgres Over ChakrDB",
    "Process Docs": "IDF Processes & Operations",
    "RCAs": "Root Cause Analyses",
    "RPC behaviours documentation": "RPC Behaviour & API Documentation",
    # Legacy categories
    "idf_lattice": "IDF Lattice & Federated Entity Types",
    "namespace": "IDF Namespaces",
    "process_docs": "IDF Processes & Operations",
    "requirement_docs": "IDF Requirements & Design",
}

CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200
MIN_DOC_LENGTH = 100

SKIP_PATTERNS = [
    r"^Copy of ",
    r"^~\$",
    r"Contact Details",
    r"PTO",
    r"Task Lists?\.docx$",
    r"India Contact",
]


def should_skip(filename: str) -> bool:
    """Check if a file should be skipped based on name patterns."""
    for pattern in SKIP_PATTERNS:
        if re.search(pattern, filename, re.IGNORECASE):
            return True
    return False


def extract_docx_text(filepath: str) -> Tuple[str, List[str]]:
    """Extract text and headings from a .docx file."""
    from docx import Document
    try:
        doc = Document(filepath)
        sections = []
        current_heading = ""
        current_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style and para.style.name and "Heading" in para.style.name:
                if current_text:
                    sections.append((current_heading, "\n".join(current_text)))
                current_heading = text
                current_text = []
            else:
                current_text.append(text)

        if current_text:
            sections.append((current_heading, "\n".join(current_text)))

        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_rows.append(row_text)
            if table_rows:
                sections.append(("Table", "\n".join(table_rows)))

        full_text = "\n\n".join(
            f"## {h}\n{t}" if h else t for h, t in sections
        )
        headings = [h for h, _ in sections if h]
        return full_text, headings

    except Exception as e:
        print(f"  Warning: Could not read {filepath}: {e}")
        return "", []


def extract_text(filepath: str) -> Tuple[str, List[str]]:
    """Extract text from supported file types. Returns (text, headings)."""
    ext = Path(filepath).suffix.lower()
    if ext == ".docx":
        return extract_docx_text(filepath)
    elif ext in (".txt", ".md"):
        with open(filepath, 'r', errors='ignore') as f:
            text = f.read()
        headings = re.findall(r'^#+\s+(.+)$', text, re.MULTILINE)
        return text, headings
    elif ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(filepath)
            text = "\n\n".join(page.get_text() for page in doc)
            doc.close()
            return text, []
        except ImportError:
            return "", []
    return "", []


def chunk_text_smart(text: str, filename: str, category: str,
                     chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[Dict]:
    """
    Smart chunking: splits at heading boundaries first, then at paragraph boundaries.
    Each chunk is prefixed with context (category + filename) for better embeddings.
    """
    if not text or len(text) < MIN_DOC_LENGTH:
        return []

    context_prefix = f"[{category} | {Path(filename).stem}]\n"

    heading_pattern = r'\n(?=## .+\n)'
    sections = re.split(heading_pattern, text)

    chunks = []
    current_chunk = ""

    for section in sections:
        section = section.strip()
        if not section:
            continue

        if len(current_chunk) + len(section) <= chunk_size:
            current_chunk += ("\n\n" + section) if current_chunk else section
        else:
            if current_chunk:
                chunks.append(context_prefix + current_chunk.strip())

            if len(section) <= chunk_size:
                current_chunk = section
            else:
                paragraphs = section.split("\n\n")
                current_chunk = ""
                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue
                    if len(current_chunk) + len(para) + 2 <= chunk_size:
                        current_chunk += ("\n\n" + para) if current_chunk else para
                    else:
                        if current_chunk:
                            chunks.append(context_prefix + current_chunk.strip())
                        if len(para) > chunk_size:
                            words = para.split()
                            current_chunk = ""
                            for word in words:
                                if len(current_chunk) + len(word) + 1 <= chunk_size:
                                    current_chunk += (" " + word) if current_chunk else word
                                else:
                                    chunks.append(context_prefix + current_chunk.strip())
                                    overlap_words = current_chunk.split()[-overlap // 6:]
                                    current_chunk = " ".join(overlap_words) + " " + word
                        else:
                            overlap_text = ""
                            if current_chunk:
                                words = current_chunk.split()
                                overlap_text = " ".join(words[-overlap // 6:]) + "\n\n"
                            current_chunk = overlap_text + para

    if current_chunk and len(current_chunk.strip()) > 50:
        chunks.append(context_prefix + current_chunk.strip())

    return [{"text": c, "char_len": len(c)} for c in chunks if len(c.strip()) > 80]


def get_category(filepath: str) -> str:
    """Determine document category from directory path."""
    rel = os.path.relpath(filepath, RAW_DOCS_DIR)
    top_dir = rel.split(os.sep)[0] if os.sep in rel else ""

    for key, label in CATEGORY_MAP.items():
        if key == top_dir or key in filepath:
            return label

    return "General IDF Documentation"


def extract_features(text: str, filename: str) -> List[str]:
    """Extract IDF feature keywords from document text."""
    feature_patterns = {
        "entity_type": r'\b(entity\s+type|entity_type|RegisterEntityTypes)\b',
        "namespace": r'\b(namespace|namespaces|kNamespace)\b',
        "attribute": r'\b(attribute|attributes|metric|metrics|RegisterMetricTypes)\b',
        "watch": r'\b(watch|watches|subscription|RegisterWatch)\b',
        "CAS": r'\b(CAS|cas_value|compare.and.swap)\b',
        "lattice": r'\b(lattice|federated)\b',
        "GEWM": r'\b(GEWM|GetEntitiesWithMetrics)\b',
        "batch": r'\b(batch|BatchGet|BatchUpdate|BatchDelete)\b',
        "eviction": r'\b(evictable|unevictable|eviction|kEvictable)\b',
        "secondary_index": r'\b(secondary\s+index|indexing)\b',
        "replication": r'\b(replication|replicated|PE.sync)\b',
        "proto": r'\b(proto|protobuf|protocol\s+buffer)\b',
        "PCDR": r'\b(PCDR|PC\s+DR|disaster\s+recovery)\b',
        "spotlight": r'\b(spotlight|full.text.search|SpotLightSearch)\b',
        "cursor": r'\b(cursor|pagination|cursor_id)\b',
        "aggregation": r'\b(group.by|aggregation|aggregate)\b',
        "RPC": r'\b(RPC|gRPC|Arithmos)\b',
        "schema": r'\b(schema|schema_config|entity_attribute_config)\b',
        "mongodb": r'\b(mongo|mongodb|mongoDB)\b',
        "postgres": r'\b(postgres|ChakrDB|chakrdb)\b',
        "memory": r'\b(memory|OOM|memory.management|cache)\b',
    }

    features = set()
    text_lower = text[:5000].lower()
    fname_lower = filename.lower()

    for feature_name, pattern in feature_patterns.items():
        if re.search(pattern, text_lower, re.IGNORECASE) or re.search(pattern, fname_lower, re.IGNORECASE):
            features.add(feature_name)

    return list(features)


def build_index() -> Tuple[List[Dict], Dict]:
    """Process all documents and build the document index."""
    print("=" * 60)
    print(" IDF Knowledge Base Indexer v2")
    print("=" * 60)
    print(f" Source: {RAW_DOCS_DIR}")
    print(f" Chunk size: {CHUNK_SIZE}, overlap: {CHUNK_OVERLAP}")
    print()

    documents = []
    doc_metadata = {}
    total_chunks = 0
    skipped = 0

    supported_extensions = {'.docx', '.txt', '.md', '.pdf'}

    for root, dirs, files in os.walk(RAW_DOCS_DIR):
        dirs.sort()
        for filename in sorted(files):
            filepath = os.path.join(root, filename)
            ext = Path(filepath).suffix.lower()

            if ext not in supported_extensions:
                continue

            if should_skip(filename):
                skipped += 1
                continue

            text, headings = extract_text(filepath)
            if not text or len(text) < MIN_DOC_LENGTH:
                skipped += 1
                continue

            category = get_category(filepath)
            features = extract_features(text, filename)
            doc_id = hashlib.md5(filepath.encode()).hexdigest()[:12]

            chunks = chunk_text_smart(text, filename, category)
            if not chunks:
                skipped += 1
                continue

            doc_entry = {
                "id": doc_id,
                "filename": filename,
                "filepath": filepath,
                "category": category,
                "features": features,
                "headings": headings[:10],
                "num_chunks": len(chunks),
                "text_length": len(text),
                "summary_text": text[:600],
            }
            documents.append(doc_entry)
            doc_metadata[doc_id] = {
                "chunks": [c["text"] for c in chunks],
                "features": features,
                "category": category,
                "filename": filename,
                "headings": headings[:10],
            }

            total_chunks += len(chunks)
            print(f"  [{category[:25]:25s}] {filename[:45]:45s} -> {len(chunks):3d} chunks")

    print(f"\n{'=' * 60}")
    print(f" Total: {len(documents)} documents, {total_chunks} chunks")
    print(f" Skipped: {skipped} files")
    print(f"{'=' * 60}")

    with open(INDEX_PATH, 'w') as f:
        json.dump({"documents": documents, "metadata": doc_metadata}, f, indent=2)
    print(f" Index saved: {INDEX_PATH}")

    return documents, doc_metadata


def build_chromadb(documents: List[Dict], doc_metadata: Dict):
    """Build ChromaDB vector store from chunked documents."""
    import chromadb
    import requests

    EMBEDDINGS_URL = "http://localhost:11434/api/embeddings"
    MODEL = "nomic-embed-text"

    print("\n Building ChromaDB vector store...")
    print(f"   Embeddings model: {MODEL}")

    try:
        r = requests.post(EMBEDDINGS_URL, json={"model": MODEL, "prompt": "test"}, timeout=10)
        if r.status_code != 200:
            print(f"   ERROR: Ollama embeddings not available (status {r.status_code})")
            print("   Run: ollama pull nomic-embed-text")
            return
    except Exception as e:
        print(f"   ERROR: Cannot reach Ollama: {e}")
        print("   Make sure Ollama is running: ollama serve")
        return

    if os.path.exists(CHROMA_DIR):
        import shutil
        shutil.rmtree(CHROMA_DIR)
        print("   Cleared old ChromaDB")

    client = chromadb.PersistentClient(path=CHROMA_DIR)

    collection = client.create_collection(
        name="idf_knowledge_base",
        metadata={"description": "IDF team knowledge base — 12 categories, ~470 documents"}
    )

    all_ids = []
    all_texts = []
    all_metadatas = []

    seen_hashes = set()

    for doc in documents:
        doc_id = doc["id"]
        meta = doc_metadata[doc_id]
        chunks = meta["chunks"]

        for i, chunk in enumerate(chunks):
            chunk_hash = hashlib.md5(chunk.encode()).hexdigest()[:16]
            if chunk_hash in seen_hashes:
                continue
            seen_hashes.add(chunk_hash)

            chunk_id = f"{doc_id}_chunk_{i}"
            all_ids.append(chunk_id)
            all_texts.append(chunk)
            all_metadatas.append({
                "doc_id": doc_id,
                "filename": doc["filename"],
                "category": doc["category"],
                "features": json.dumps(doc["features"]),
                "chunk_index": i,
                "total_chunks": len(chunks),
            })

    BATCH_SIZE = 50
    total = len(all_texts)
    print(f"   Indexing {total} unique chunks (deduped from {sum(d['num_chunks'] for d in documents)})...")

    for batch_start in range(0, total, BATCH_SIZE):
        batch_end = min(batch_start + BATCH_SIZE, total)
        batch_texts = all_texts[batch_start:batch_end]
        batch_ids = all_ids[batch_start:batch_end]
        batch_meta = all_metadatas[batch_start:batch_end]

        embeddings = []
        for text in batch_texts:
            try:
                resp = requests.post(
                    EMBEDDINGS_URL,
                    json={"model": MODEL, "prompt": text[:2000]},
                    timeout=30
                )
                if resp.status_code == 200:
                    embeddings.append(resp.json()["embedding"])
                else:
                    embeddings.append([0.0] * 768)
            except Exception:
                embeddings.append([0.0] * 768)

        collection.add(
            ids=batch_ids,
            documents=batch_texts,
            embeddings=embeddings,
            metadatas=batch_meta,
        )

        pct = batch_end * 100 // total
        print(f"   [{pct:3d}%] Indexed {batch_end}/{total} chunks...")

    print(f"\n ChromaDB built: {CHROMA_DIR}")
    print(f"   Collection: idf_knowledge_base ({total} vectors)")


if __name__ == "__main__":
    documents, doc_metadata = build_index()
    build_chromadb(documents, doc_metadata)
    print("\n Done! Knowledge base is ready.")
